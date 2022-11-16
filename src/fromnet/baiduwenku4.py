#百度文库采集
#20200803微信：huguo00289
#https://wenku.baidu.com/view/312ce9da0129bd64783e0912a216147916117e27.html
# -*- coding: UTF-8 -*-

import requests,re
from lxml import etree
from docx import Document

def get_detail(url):
    #url = 'https://wenku.baidu.com/view/312ce9da0129bd64783e0912a216147916117e27.html'
    header = {'User-agent': 'Googlebot'}
    response = requests.get(url , headers = header).content.decode('utf-8')
    #print(response)
    title_ze=r'<title>(.+?)_百度文库</title>'
    div_ze=r'<div class="bd doc-reader">(.+?)<div class="aside">'
    title=re.findall(title_ze,response,re.S)[0]
    div=re.findall(div_ze,response,re.S)[0]
    div=etree.HTML(div)
    details=div.xpath('//div//text()')
    #detail='\n'.join(details)
    data=title,details
    print(data)
    return data



def get_word(data):
    document = Document()
    document.add_heading(data[0])

    for detail in data[1]:
        document.add_paragraph(detail) #添加段落


    document.save(f'{data[0]}.docx')

if __name__=='__main__':
    url="https://wenku.baidu.com/view/b12ff5a18462caaedd3383c4bb4cf7ec4afeb69e?fr=sogou&_wkts_=1668521533825"
    text=get_detail(url)
    get_word(text)